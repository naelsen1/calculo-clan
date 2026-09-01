import streamlit as st
import pandas as pd
import openpyxl
import os
import uuid
import sys
import subprocess
from datetime import datetime, date
from io import BytesIO

# Manipulação de Word para a Notificação
try:
    import docx
    DOCX_DISPONIVEL = True
except ImportError:
    DOCX_DISPONIVEL = False

# Importação para conversão em Windows via COM API
try:
    import win32com.client as win32
    import pythoncom
    EXCEL_DISPONIVEL = True
except ImportError:
    EXCEL_DISPONIVEL = False

st.set_page_config(
    page_title="Calculadora CNR - Inciso V",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 📌 DICIONÁRIO DE MESES EM PORTUGUÊS (GARANTE TRADUÇÃO NO STREAMLIT CLOUD / LINUX)
MESES_PT = {
    1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
    5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
    9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
}

def obter_data_extenso_pt(dt=None):
    """Retorna a data atual ou informada por extenso em português brasileiro."""
    if dt is None:
        dt = datetime.now()
    nome_mes = MESES_PT.get(dt.month, "")
    return f"{dt.day} de {nome_mes} de {dt.year}"

def fmt_br(valor, decimais=2):
    if pd.isnull(valor):
        return "0,00"
    formato = f"{{:,.{decimais}f}}"
    return formato.format(valor).replace(",", "X").replace(".", ",").replace("X", ".")

def carregar_arquivo(uploaded_file):
    if uploaded_file.name.endswith(('xlsx', 'xls')):
        return pd.read_excel(uploaded_file)
    else:
        encodings = ['latin1', 'iso-8859-1', 'cp1252', 'utf-8']
        for enc in encodings:
            try:
                uploaded_file.seek(0)
                return pd.read_csv(uploaded_file, sep=None, engine='python', encoding=enc)
            except (UnicodeDecodeError, Exception):
                continue
        raise Exception("Não foi possível carregar o arquivo CSV.")

def converter_para_pdf_universal(caminho_entrada, caminho_saida):
    """
    Função Híbrida:
    - Se for Windows com Office: mantém a execução original via win32com.
    - Se for Linux / Streamlit Cloud: utiliza LibreOffice em modo headless.
    """
    if sys.platform.startswith("win") and EXCEL_DISPONIVEL:
        try:
            pythoncom.CoInitialize()
            if caminho_entrada.endswith('.xlsx'):
                excel_app = win32.Dispatch("Excel.Application")
                try: excel_app.Visible = False
                except Exception: pass
                excel_app.DisplayAlerts = False
                
                wb_com = excel_app.Workbooks.Open(caminho_entrada)
                ws_com = wb_com.Worksheets(1)
                ws_com.PageSetup.Zoom = False
                ws_com.PageSetup.FitToPagesWide = 1
                ws_com.PageSetup.FitToPagesTall = 1
                
                wb_com.ExportAsFixedFormat(0, caminho_saida)
                wb_com.Close(False)
                excel_app.Quit()
            elif caminho_entrada.endswith('.docx'):
                word_app = win32.Dispatch("Word.Application")
                try: word_app.Visible = False
                except Exception: pass
                
                doc_com = word_app.Documents.Open(caminho_entrada)
                doc_com.SaveAs(caminho_saida, FileFormat=17) # 17 = wdFormatPDF
                doc_com.Close(False)
                word_app.Quit()
            return os.path.exists(caminho_saida)
        except Exception as e:
            st.error(f"Erro na conversão via Office (Windows): {e}")
            return False
        finally:
            pythoncom.CoUninitialize()
    else:
        try:
            diretorio_saida = os.path.dirname(caminho_saida)
            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", diretorio_saida,
                caminho_entrada
            ]
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            
            pdf_gerado_temp = os.path.splitext(caminho_entrada)[0] + ".pdf"
            if os.path.exists(pdf_gerado_temp) and pdf_gerado_temp != caminho_saida:
                os.rename(pdf_gerado_temp, caminho_saida)
                
            return os.path.exists(caminho_saida)
        except Exception as e:
            st.error(f"Erro na conversão via LibreOffice (Nuvem): {e}")
            return False

def preencher_modelo_excel(caminho_modelo, dados_prod, maior_ciclo_row, consumo_diario, dias_cobranca, usou_minima, dt_ini_efetiva, dt_fim_efetiva):
    wb = openpyxl.load_workbook(caminho_modelo)
    ws = wb['DIARIO'] if 'DIARIO' in wb.sheetnames else wb.active
    
    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    
    # DATAS E NOME POR EXTENSO NO CABEÇALHO DO EXCEL (GARANTE PT-BR)
    ws['A6'] = f"Goiânia, {obter_data_extenso_pt()}"
    
    ws['B8'] = str(dados_prod.get('CTA_NOME', 'N/A'))
    ws['B9'] = str(dados_prod.get('UC', 'N/A'))
    ws['B10'] = str(dados_prod.get('TOI', 'N/A'))
    
    dt_ini_str = dt_ini_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_ini_efetiva) else ''
    dt_fim_str = dt_fim_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_fim_efetiva) else ''
    
    ws['D15'] = dt_ini_str
    ws['F15'] = dt_fim_str
    ws['H15'] = dias_cobranca
    
    ws['A19'] = "Consumo Base Dia (kWh/dia)"
    ws['I19'] = round(consumo_diario, 4)
    ws['I20'] = f"=I19*H15"
    ws['I21'] = 0
    ws['I22'] = f"=I20-I21"
    
    dt_hoje = datetime.now()
    ws['A25'] = f"d) Tarifa atual ref.: {dt_hoje.strftime('%m/%Y')}"
    
    criterio_str = "MÍNIMA" if usou_minima else "LIDA"
    mes_ref = int(maior_ciclo_row['MES_REF'].values[0])
    ano_ref = int(maior_ciclo_row['ANO_REF'].values[0])
    kwh_ref = maior_ciclo_row['CONSUMOLIDOKWH'].values[0]
    
    ws['A32'] = f"Período da base de cálculo (Maior Pós - {criterio_str}): Ref. {mes_ref:02d}/{ano_ref} com {kwh_ref:,.0f} kWh."
    ws['A35'] = f"Referência Período: {dt_ini_str} a {dt_fim_str} ({dias_cobranca} dias) - Ciclo de Referência {mes_ref:02d}/{ano_ref} - {kwh_ref:,.0f} kWh / 30 = {consumo_diario:,.4f} kWh"
    
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def gerar_e_converter_pdf_demanda(caminho_modelo, dados_prod, maior_ciclo_row, consumo_diario, dias_cobranca, usou_minima, uc_nome, dt_ini_efetiva, dt_fim_efetiva):
    uid = uuid.uuid4().hex[:8]
    temp_excel = os.path.abspath(f"temp_carta_{uc_nome}_{uid}.xlsx")
    temp_pdf = os.path.abspath(f"temp_carta_{uc_nome}_{uid}.pdf")
    
    wb = openpyxl.load_workbook(caminho_modelo)
    ws = wb['DIARIO'] if 'DIARIO' in wb.sheetnames else wb.active
    
    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    
    ws['A6'] = f"Goiânia, {obter_data_extenso_pt()}"
    ws['B8'] = str(dados_prod.get('CTA_NOME', 'N/A'))
    ws['B9'] = str(dados_prod.get('UC', 'N/A'))
    ws['B10'] = str(dados_prod.get('TOI', 'N/A'))
    
    dt_ini_str = dt_ini_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_ini_efetiva) else ''
    dt_fim_str = dt_fim_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_fim_efetiva) else ''
    
    ws['D15'] = dt_ini_str
    ws['F15'] = dt_fim_str
    ws['H15'] = dias_cobranca
    
    ws['A19'] = "Consumo Base Dia (kWh/dia)"
    ws['I19'] = round(consumo_diario, 4)
    ws['I20'] = f"=I19*H15"
    ws['I21'] = 0
    ws['I22'] = f"=I20-I21"
    
    dt_hoje = datetime.now()
    ws['A25'] = f"d) Tarifa atual ref.: {dt_hoje.strftime('%m/%Y')}"
    
    criterio_str = "MÍNIMA" if usou_minima else "LIDA"
    mes_ref = int(maior_ciclo_row['MES_REF'].values[0])
    ano_ref = int(maior_ciclo_row['ANO_REF'].values[0])
    kwh_ref = maior_ciclo_row['CONSUMOLIDOKWH'].values[0]
    
    ws['A32'] = f"Período da base de cálculo (Maior Pós - {criterio_str}): Ref. {mes_ref:02d}/{ano_ref} com {kwh_ref:,.0f} kWh."
    ws['A35'] = f"Referência Período: {dt_ini_str} a {dt_fim_str} ({dias_cobranca} dias) - Ciclo de Referência {mes_ref:02d}/{ano_ref} - {kwh_ref:,.0f} kWh / 30 = {consumo_diario:,.4f} kWh"
    
    wb.save(temp_excel)
    
    pdf_bytes = None
    if converter_para_pdf_universal(temp_excel, temp_pdf):
        with open(temp_pdf, "rb") as f:
            pdf_bytes = f.read()
            
    for temp_file in [temp_excel, temp_pdf]:
        if os.path.exists(temp_file):
            try: os.remove(temp_file)
            except Exception: pass
            
    return pdf_bytes

def gerar_notificacao_pdf(caminho_docx, dic_substituicoes, uc_nome):
    uid = uuid.uuid4().hex[:8]
    temp_docx = os.path.abspath(f"temp_notif_{uc_nome}_{uid}.docx")
    temp_pdf = os.path.abspath(f"temp_notif_{uc_nome}_{uid}.pdf")
    
    doc = docx.Document(caminho_docx)
    
    def aplicar_substituicoes_em_paragrafo(p):
        texto_original = p.text
        texto_modificado = texto_original
        for k, v in dic_substituicoes.items():
            if k in texto_modificado:
                texto_modificado = texto_modificado.replace(k, str(v))
        
        if texto_modificado != texto_original:
            p.text = texto_modificado

    for p in doc.paragraphs:
        aplicar_substituicoes_em_paragrafo(p)
        
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    aplicar_substituicoes_em_paragrafo(p)
                    
    for section in doc.sections:
        for header_p in section.header.paragraphs:
            aplicar_substituicoes_em_paragrafo(header_p)
        for footer_p in section.footer.paragraphs:
            aplicar_substituicoes_em_paragrafo(footer_p)

    doc.save(temp_docx)
    
    pdf_bytes = None
    if converter_para_pdf_universal(temp_docx, temp_pdf):
        with open(temp_pdf, "rb") as f:
            pdf_bytes = f.read()
            
    for temp_file in [temp_docx, temp_pdf]:
        if os.path.exists(temp_file):
            try: os.remove(temp_file)
            except Exception: pass
            
    return pdf_bytes

st.title("Calculadora de CNR Clandestino (Art. 595, Inciso V)")
st.write("Análise individual por UC e TOI com base no maior consumo pós-regularização.")

# BARRA LATERAL (SIDEBAR)
st.sidebar.header("Upload dos Arquivos")
file_historico = st.sidebar.file_uploader("1. Upload Planilha 1 (Histórico UC - Obrigatório)", type=['xlsx', 'xls', 'csv'])
file_producao = st.sidebar.file_uploader("2. Upload Planilha 2 (Base de Produção TR - Opcional)", type=['xlsx', 'xls', 'csv'])

st.sidebar.markdown("---")
st.sidebar.header("Configuração de Tarifa")
tarifa_vigente_input = st.sidebar.number_input(
    "Tarifa Vigente (R$/kWh):",
    value=1.17466,
    format="%.5f",
    step=0.00001,
    help="Tarifa aplicável para a apuração do valor total faturado."
)

if file_historico:
    try:
        df_hist = carregar_arquivo(file_historico)
        df_hist['UC'] = df_hist['UC'].astype(str).str.strip().str.replace(r'\.0$', '', regex=True)
        
        df_prod = None
        if file_producao:
            try:
                df_prod = carregar_arquivo(file_producao)
                df_prod['UC'] = df_prod['UC'].astype(str).str.strip().str.replace(r'\.0$', '', regex=True)
            except Exception as e:
                st.sidebar.warning(f"⚠️ Não foi possível ler a Planilha 2. Entrando em Modo Contingência.")
    except Exception as e:
        st.error(f"Erro ao carregar o Histórico (Planilha 1): {e}")
        st.stop()

    ucs_hist = list(df_hist['UC'].dropna().unique())
    
    if not ucs_hist:
        st.error("Nenhuma UC encontrada na Planilha de Histórico.")
    else:
        ucs_comuns = []
        if df_prod is not None:
            ucs_prod = set(df_prod['UC'].dropna().unique())
            ucs_comuns = list(set(ucs_hist).intersection(ucs_prod))

        if ucs_comuns:
            opcoes_uc = {}
            for uc in ucs_comuns:
                toi_val = df_prod[df_prod['UC'] == uc]['TOI'].values
                toi_str = str(toi_val[0]) if len(toi_val) > 0 else "S/N"
                opcoes_uc[f"UC: {uc} | TOI: {toi_str}"] = uc
            st.success(f"✓ {len(ucs_comuns)} UC(s) em comum identificadas entre o Histórico e a Base de Produção.")
        else:
            if df_prod is not None:
                st.warning("⚠️ Não foram encontradas UCs em comum entre a Planilha 1 e a Planilha 2. Modo Contingência ativado.")
            else:
                st.info("ℹ️ Modo Contingência Ativo: Apurando histórico diretamente da Planilha 1.")
            
            opcoes_uc = {f"UC: {uc} (Histórico)": uc for uc in ucs_hist}

        selected_label = st.selectbox("Selecione a UC para analisar o cálculo detalhado", list(opcoes_uc.keys()))
        selected_uc = opcoes_uc[selected_label]
        
        if selected_uc:
            if df_prod is not None and selected_uc in df_prod['UC'].values:
                dados_uc_prod = df_prod[df_prod['UC'] == selected_uc].iloc[0].to_dict()
            else:
                dados_uc_prod = {}
                
            valor_coluna_d = dados_uc_prod.get(list(dados_uc_prod.keys())[3], None) if len(dados_uc_prod) > 3 else None
            data_inspecao_coluna_d = pd.to_datetime(valor_coluna_d, errors='coerce') if valor_coluna_d else None
            
            # 1. DADOS DO CLIENTE E DO CÁLCULO
            st.subheader("1. Dados do Cliente e do Cálculo")
            
            col_man1, col_man2 = st.columns(2)
            
            dt_declarada_input = col_man1.date_input(
                "Início Declarado pelo Cliente (Opcional):",
                value=None,
                format="DD/MM/YYYY"
            )
            
            dt_inspecao_tr_input = col_man2.date_input(
                "Data de Inspeção / Normalização (TR):",
                value=data_inspecao_coluna_d.date() if pd.notnull(data_inspecao_coluna_d) else None,
                format="DD/MM/YYYY"
            )

            if dt_inspecao_tr_input is not None:
                data_inspecao_usada = pd.Timestamp(dt_inspecao_tr_input)
                dt_fim_efetiva = data_inspecao_usada
            elif pd.notnull(data_inspecao_coluna_d):
                data_inspecao_usada = data_inspecao_coluna_d
                dt_fim_efetiva = pd.to_datetime(dados_uc_prod.get('DATA_FINAL_DIA', data_inspecao_coluna_d), errors='coerce')
            else:
                data_inspecao_usada = pd.Timestamp(datetime.now().date())
                dt_fim_efetiva = data_inspecao_usada

            data_insp_str = data_inspecao_usada.strftime('%d/%m/%Y') if pd.notnull(data_inspecao_usada) else 'N/A'

            # BUSCA GARANTIDA DO NOME DO TITULAR MAIS RECENTE/ATUAL DA UC
            nome_cliente_padrao = "N/A"
            dados_uc_hist_busca = df_hist[df_hist['UC'] == selected_uc].copy()
            
            if 'NOMECLIENTE' in dados_uc_hist_busca.columns and len(dados_uc_hist_busca) > 0:
                dados_uc_hist_busca['DT_FIM_PARSED'] = pd.to_datetime(dados_uc_hist_busca['DATA_FINAL'], errors='coerce')
                
                pos_insp = dados_uc_hist_busca[dados_uc_hist_busca['DT_FIM_PARSED'] > data_inspecao_usada].sort_values('DT_FIM_PARSED')
                if len(pos_insp) > 0 and pd.notnull(pos_insp['NOMECLIENTE'].values[0]):
                    nome_cliente_padrao = str(pos_insp['NOMECLIENTE'].values[0]).strip()
                else:
                    hist_ordenado = dados_uc_hist_busca.sort_values('DT_FIM_PARSED', ascending=False)
                    nomes_validos = hist_ordenado['NOMECLIENTE'].dropna().unique()
                    if len(nomes_validos) > 0:
                        nome_cliente_padrao = str(nomes_validos[0]).strip()

            if nome_cliente_padrao == "N/A" and 'CTA_NOME' in dados_uc_prod and pd.notnull(dados_uc_prod['CTA_NOME']):
                nome_cliente_padrao = str(dados_uc_prod['CTA_NOME'])

            # CAMPOS CADASTRAIS
            st.markdown("**Campos Cadastrais:**")
            col_cad1, col_cad2, col_cad3 = st.columns(3)
            
            nome_cliente_input = col_cad1.text_input("Nome do Cliente Atual:", value=nome_cliente_padrao)
            toi_input = col_cad2.text_input("Número do TOI / Inspeção:", value=str(dados_uc_prod.get('TOI', 'SN')))
            endereco_input = col_cad3.text_input("Endereço do Cliente:", value=str(dados_uc_prod.get('ENDERECO', dados_uc_prod.get('LOGRADOURO', 'N/A'))))

            dados_uc_prod['CTA_NOME'] = nome_cliente_input
            dados_uc_prod['UC'] = selected_uc
            dados_uc_prod['TOI'] = toi_input

            dt_ini_cob = pd.to_datetime(dados_uc_prod.get('DATA_INICIAL_DIA', None), errors='coerce')
            
            if dt_declarada_input is not None:
                dt_ini_efetiva = pd.Timestamp(dt_declarada_input)
                if pd.notnull(dt_fim_efetiva):
                    dias_calculados = (dt_fim_efetiva - dt_ini_efetiva).days
                    if dias_calculados > 180:
                        st.warning(f"⚠️ **Alerta:** O período calculado ({dias_calculados} dias) ultrapassa o limite de **180 dias**. "
                                   f"Período limitado a 180 dias. **Início Sugerido:** `{(dt_fim_efetiva - pd.Timedelta(days=180)).strftime('%d/%m/%Y')}`.")
                        dias_cobranca = 180
                        dt_ini_efetiva = dt_fim_efetiva - pd.Timedelta(days=180)
                    else:
                        dias_cobranca = max(0, dias_calculados)
                else:
                    dias_cobranca = 0
            elif dt_inspecao_tr_input is not None or pd.isnull(dt_ini_cob):
                dias_cobranca = 180
                dt_ini_efetiva = data_inspecao_usada - pd.Timedelta(days=180)
                st.info(f"ℹ️ **Período Retroativo Automático:** Calculando 180 dias para trás a partir da Data de Normalização/Inspeção ({data_inspecao_usada.strftime('%d/%m/%Y')}).")
            else:
                dt_ini_efetiva = dt_ini_cob
                dt_fim_cob = pd.to_datetime(dados_uc_prod.get('DATA_FINAL_DIA', None), errors='coerce')
                if pd.notnull(dados_uc_prod.get('PERIODODIAS')) and int(dados_uc_prod.get('PERIODODIAS')) > 0:
                    dias_cobranca = int(dados_uc_prod.get('PERIODODIAS'))
                elif pd.notnull(dt_ini_cob) and pd.notnull(dt_fim_cob):
                    dias_cobranca = (dt_fim_cob - dt_ini_cob).days
                else:
                    dias_cobranca = 180
            
            st.write(f"**Resumo do Período Efetivo:** `{dias_cobranca} Dia(s)` | **Início:** `{dt_ini_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_ini_efetiva) else 'N/A'}` | **Fim:** `{dt_fim_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_fim_efetiva) else 'N/A'}`")
            
            st.markdown("---")
            st.subheader("2. Histórico de Pós-Faturamento")
            
            dados_uc_hist = df_hist[df_hist['UC'] == selected_uc].copy()
            dados_uc_hist['DATA_FINAL'] = pd.to_datetime(dados_uc_hist['DATA_FINAL'], errors='coerce')
            
            exibir_todo_historico = st.checkbox("🔍 Exibir histórico completo (incluindo ciclos anteriores à inspeção)", value=False)
            
            if exibir_todo_historico:
                st.write("**Histórico Completo da UC (Ordenado por Mês 1-12):**")
                df_exibicao = dados_uc_hist.copy()
                df_exibicao['MES_REF'] = pd.to_numeric(df_exibicao['MES_REF'], errors='coerce')
                df_exibicao = df_exibicao.sort_values(by=['ANO_REF', 'MES_REF'])
                df_exibicao['DATA_INICIAL'] = df_exibicao['DATA_INICIAL'].astype(str)
                df_exibicao['DATA_FINAL'] = df_exibicao['DATA_FINAL'].dt.strftime('%d/%m/%Y')
                
                cols_exib = [c for c in ['ANO_REF', 'MES_REF', 'DATA_INICIAL', 'DATA_FINAL', 'CONSUMOLIDOKWH', 'CRITERIOFATURAMENTO', 'NOMECLIENTE'] if c in df_exibicao.columns]
                st.dataframe(df_exibicao[cols_exib], use_container_width=True)
            
            pos_faturamento = dados_uc_hist[dados_uc_hist['DATA_FINAL'] > data_inspecao_usada].copy() if pd.notnull(data_inspecao_usada) else pd.DataFrame()
            
            if len(pos_faturamento) > 1:
                pos_faturamento = pos_faturamento.sort_values('DATA_FINAL')
                
                ciclos_apos_descarto = pos_faturamento.iloc[1:]
                ciclos_validos = ciclos_apos_descarto.head(3).copy()
                
                pos_lidas = ciclos_validos[ciclos_validos['CRITERIOFATURAMENTO'].str.upper() == 'LIDA'] if 'CRITERIOFATURAMENTO' in ciclos_validos.columns else pd.DataFrame()
                
                usou_minima = False
                if len(pos_lidas) > 0:
                    base_selecionada = pos_lidas
                else:
                    pos_minima = ciclos_validos[ciclos_validos['CRITERIOFATURAMENTO'].str.upper() == 'MÍNIMA'] if 'CRITERIOFATURAMENTO' in ciclos_validos.columns else pd.DataFrame()
                    if len(pos_minima) > 0:
                        base_selecionada = pos_minima
                        usou_minima = True
                    else:
                        base_selecionada = ciclos_validos
                
                maior_ciclo_row = base_selecionada.nlargest(1, 'CONSUMOLIDOKWH')
                maior_kwh = maior_ciclo_row['CONSUMOLIDOKWH'].values[0]
                
                st.info(f"Busca realizada após a inspeção de **{data_insp_str}**. Analisados os **3 ciclos subsequentes** ao 1º de expurgo.")
                
                if usou_minima:
                    st.warning("⚠️ **Atenção:** Não foram encontrados ciclos com faturamento 'LIDA' na janela de 3 ciclos. O cálculo utilizou o maior ciclo 'MÍNIMA'.")
                
                st.write("**Janela dos 3 Ciclos Pós Observados:**")
                df_janela_exib = ciclos_validos.copy()
                df_janela_exib['DATA_INICIAL'] = df_janela_exib['DATA_INICIAL'].astype(str)
                df_janela_exib['DATA_FINAL'] = df_janela_exib['DATA_FINAL'].dt.strftime('%d/%m/%Y')
                
                cols_janela = [c for c in ['ANO_REF', 'MES_REF', 'DATA_INICIAL', 'DATA_FINAL', 'CONSUMOLIDOKWH', 'CRITERIOFATURAMENTO', 'NOMECLIENTE'] if c in df_janela_exib.columns]
                st.dataframe(df_janela_exib[cols_janela], use_container_width=True)
                
                st.write("**Maior Consumo Pós-Faturamento Selecionado para Cálculo:**")
                df_maior_exib = maior_ciclo_row.copy()
                df_maior_exib['DATA_INICIAL'] = df_maior_exib['DATA_INICIAL'].astype(str)
                df_maior_exib['DATA_FINAL'] = df_maior_exib['DATA_FINAL'].dt.strftime('%d/%m/%Y')
                
                cols_maior = [c for c in ['ANO_REF', 'MES_REF', 'DATA_INICIAL', 'DATA_FINAL', 'CONSUMOLIDOKWH', 'CRITERIOFATURAMENTO', 'NOMECLIENTE'] if c in df_maior_exib.columns]
                st.dataframe(df_maior_exib[cols_maior], use_container_width=True)
                
                st.markdown("---")
                st.subheader("3. Memória de Cálculo Estimada")
                
                consumo_diario = maior_kwh / 30.0
                consumo_estimado_total = consumo_diario * dias_cobranca
                
                valor_total_debito = consumo_estimado_total * tarifa_vigente_input
                
                col_calc1, col_calc2, col_calc3, col_calc4 = st.columns(4)
                col_calc1.metric(label="Maior Consumo Pós (3m)", value=f"{fmt_br(maior_kwh, 2)} kWh")
                col_calc2.metric(label="Consumo Base Diário", value=f"{fmt_br(consumo_diario, 2)} kWh/dia")
                col_calc3.metric(label=f"Estimado Total ({dias_cobranca}d)", value=f"{fmt_br(consumo_estimado_total, 2)} kWh")
                col_calc4.metric(label=f"Valor Total (Tarifa R$ {tarifa_vigente_input:.5f})", value=f"R$ {fmt_br(valor_total_debito, 2)}")
                
                st.write("**Clique no ícone à direita para copiar o Consumo Diário:**")
                st.code(f"{consumo_diario:.4f}".replace(".", ","), language="text")
                
                st.markdown("---")
                st.subheader("4. Emissão dos Documentos Oficiais")
                
                caminho_modelo_excel = "Modelo_Carta_Lesao_Limpo.xlsx"
                caminho_modelo_word = "Modelo_Notificacao_Oficial.docx"
                
                toi_str = str(toi_input)
                
                col_doc1, col_doc2 = st.columns(2)
                
                # DOCUMENTO 1: DETALHAMENTO DO FATURAMENTO (EXCEL / PDF)
                with col_doc1:
                    st.markdown("### 📊 Detalhamento do Faturamento")
                    if os.path.exists(caminho_modelo_excel):
                        excel_preenchido = preencher_modelo_excel(
                            caminho_modelo_excel, dados_uc_prod, maior_ciclo_row, consumo_diario, dias_cobranca, usou_minima, dt_ini_efetiva, dt_fim_efetiva
                        )
                        st.download_button(
                            label="Baixar Detalhamento em Excel (.xlsx)",
                            data=excel_preenchido,
                            file_name=f"Detalhamento_do_Faturamento_UC_{selected_uc}_TOI_{toi_str}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        
                        if st.button("Gerar Detalhamento em PDF (.pdf)", type="primary", use_container_width=True):
                            with st.spinner("Gerando PDF do Detalhamento..."):
                                pdf_detalhamento = gerar_e_converter_pdf_demanda(
                                    caminho_modelo_excel, dados_uc_prod, maior_ciclo_row, consumo_diario, dias_cobranca, usou_minima, selected_uc, dt_ini_efetiva, dt_fim_efetiva
                                )
                                if pdf_detalhamento:
                                    st.download_button(
                                        label="💾 Salvar PDF do Detalhamento",
                                        data=pdf_detalhamento,
                                        file_name=f"Detalhamento_do_Faturamento_UC_{selected_uc}_TOI_{toi_str}.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                    else:
                        st.error(f"⚠️ O arquivo `{caminho_modelo_excel}` não foi encontrado.")
                
                # DOCUMENTO 2: NOTIFICAÇÃO DE IRREGULARIDADE (CARTA NOTIFICAÇÃO GOIÁS)
                with col_doc2:
                    st.markdown("### 📄 Notificação de Irregularidade")
                    if os.path.exists(caminho_modelo_word) and DOCX_DISPONIVEL:
                        # 📌 DATA DE EMISSÃO FORMATADA EM PORTUGUÊS
                        dt_hoje_str = datetime.now().strftime('%d.%m.%Y')
                        dt_ini_str = dt_ini_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_ini_efetiva) else ''
                        dt_fim_str = dt_fim_efetiva.strftime('%d/%m/%Y') if pd.notnull(dt_fim_efetiva) else ''
                        
                        dic_substituicoes = {
                            "{{ NUMERO_INSPECAO }}": toi_str,
                            "{{ DATA_EMISSAO }}": dt_hoje_str,
                            "{{ NOME_CLIENTE }}": nome_cliente_input,
                            "{{ ENDERECO_CLIENTE }}": endereco_input if (endereco_input and endereco_input.strip() != "N/A") else "Endereço não cadastrado",
                            "{{ UC_INSTALACAO }}": selected_uc,
                            "{{ DATA_INSPECAO }}": data_insp_str,
                            "{{ PERIODO_IRREGULARIDADE }}": dt_ini_str,
                            "{{ CONSUMO_REGISTRADO }}": "0",
                            "{{ CONSUMO_APURADO }}": fmt_br(consumo_estimado_total, 0),
                            "{{ VALOR_TOTAL }}": fmt_br(valor_total_debito, 2),
                            "{{ JUSTIFICATIVA_CRITERIO }}": "Não foi possível, no momento da inspeção, levantar a carga total da unidade consumidora, ou, pelo menos, identificar quais equipamentos estavam sendo alimentados pelo desvio de energia."
                        }
                        
                        if st.button("Gerar Notificação Oficial em PDF (.pdf)", type="primary", use_container_width=True):
                            with st.spinner("Gerando Notificação Oficial de Goiás em PDF..."):
                                pdf_notificacao = gerar_notificacao_pdf(caminho_modelo_word, dic_substituicoes, selected_uc)
                                if pdf_notificacao:
                                    st.download_button(
                                        label="💾 Salvar Notificação Oficial (PDF)",
                                        data=pdf_notificacao,
                                        file_name=f"Notificacao_Irregularidade_UC_{selected_uc}_TOI_{toi_str}.pdf",
                                        mime="application/pdf",
                                        use_container_width=True
                                    )
                    else:
                        st.error(f"⚠️ O modelo Word `{caminho_modelo_word}` não foi encontrado.")
            else:
                st.warning(f"Apenas {len(pos_faturamento)} ciclo(s) encontrados após {data_insp_str}. Mínimo necessário: 2 ciclos pós-inspeção.")
else:
    st.info("Por favor, faça o upload da Planilha 1 (Histórico UC) na barra lateral para iniciar.")